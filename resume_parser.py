# Comprehensive Resume Parser Application
# This script combines all modules to extract, parse, enhance, and reformat resumes.

# ==============================================================================
# 0. DEPENDENCY INSTALLATION
# ==============================================================================
# Before running, please install the required libraries using pip:
# pip install PyMuPDF python-docx google-generativeai

import fitz  # PyMuPDF
import docx
import re
import os
import argparse
import json
import base64
from google import genai
from google.genai import types
import docx.shared  # For handling font sizes
from docx import Document
from docx.shared import Pt # For handling font sizes footer
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from docx.oxml import OxmlElement
from dotenv import load_dotenv
from docx.opc.constants import RELATIONSHIP_TYPE as RT
load_dotenv()  # This will load variables from .env file into the environment
import string


# ==============================================================================
# 1. TEXT EXTRACTION MODULE
# ==============================================================================
# This module is responsible for extracting raw text from PDF and DOCX files.

def extract_text(file_path):
    """
    Extracts clean plain text from a resume file (PDF or DOCX).

    Args:
        file_path (str): The path to the resume file.

    Returns:
        str: The extracted raw plain text, or None if an error occurs.
    """
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return None

    file_extension = os.path.splitext(file_path)[1].lower()
    text = ""

    try:
        if file_extension == '.pdf':
            # Use PyMuPDF (fitz) to extract text from PDF
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            if not text.strip():
                print("Warning: PDF appears to be scanned or contains no text.")
        
        elif file_extension == '.docx':
            # Use python-docx to extract text from DOCX
            doc = docx.Document(file_path)
            # Extract paragraphs
            for para in doc.paragraphs:
                text += para.text + '\n'
                
                
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text += para.text + '\n'
        
        else:
            print(f"Error: Unsupported file type '{file_extension}'. Please provide a PDF or DOCX file.")
            return None
            
    except Exception as e:
        print(f"An error occurred while extracting text: {e}")
        return None

    return text.strip()

# ==============================================================================
# 2. RESUME CONTENT PARSING MODULE
# ==============================================================================
# This module parses the raw text into a structured dictionary.




def parse_content(text, api_key):
    """
    Parses raw resume text into a structured dictionary using Gemini LLM.

    Args:
        text (str): Raw extracted resume text.
        api_key (str): Gemini API key.

    Returns:
        dict: Parsed resume data.
    """
    if not api_key:
        raise ValueError("Gemini API key not set.")

    client = genai.Client(api_key=api_key)

    json_schema = {
        "name": "Full name of candidate",
        "email": "Primary email address",
        "phone": "Primary phone number",
        "linkedin": "LinkedIn URL if available",
        "github": "GitHub URL if available",
        "skills": ["List of skills"],
        "skills_matrix": ["if available in the resume",{"skills": "Skill name","years_experience": "Years of experience", "last_used": "Most recent year used", "proficiency": "Beginner/Intermediate/Advanced"}],
        "certifications": [{"name": "Certification name", "issuer": "Issuing organization"}],
        "summary": "Professional summary or objective statement or candidate strength if avaiable",
        "education": [
            {
                "degree": "Degree name",
                "major": "Major/Field of study",
                "university": "University/Institution",
                "start_date": "YYYY or MM/YYYY if available",
                "end_date": "YYYY or MM/YYYY if available",
                "gpa": "GPA if mentioned"
            }
        ],
        "experience": [
            {
                "job_title": "Title of the position",
                "company": "Employer name",
                "start_date": "YYYY or MM/YYYY",
                "end_date": "YYYY or MM/YYYY or 'Present'",
                "Description": ["Bullet point list of Description points"]
            }
        ],
        "projects": [
            {
                "date_range": "Example: YYYY-YYYY or MM/YYYY - MM/YYYY",
                "project_name": "Title of the project or heading mentioned",
                "client": "Company name or client if mentioned",
                "content": ["Content point 1.", "Point 2.", "Point 3 if available"],
                "technologies": "Technologies used, if available",
                "environment": "Environment if mentioned/available"
            }
        ],
        "awards": [
            {
                "name": "Award name",
                "issuer": "Issuing organization",
                "year": "YYYY"
            }
        ]
    }

    prompt = f"""
    You are a resume parser. Extract the following structured JSON data from the given resume text if available:

    JSON schema:
    {json.dumps(json_schema, indent=2)}

    Resume text:
    {text}

    Rules:
    - Return ONLY valid JSON.
    - Do not include any text outside the JSON.
    - If a field is missing, return an empty string or empty list as appropriate.
    - Do not invent new information; extract only what is explicitly present in the resume.
    """

    try:
        content = types.Content(role="user", parts=[types.Part.from_text(text=prompt)])
        config = types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=-1))

        response_text = ""
        for chunk in client.models.generate_content_stream(
            model="gemini-2.5-pro",
            contents=[content],
            config=config
        ):
            response_text += chunk.text
        
        # Clean up the response to ensure it's valid JSON
        cleaned_response = response_text.strip()
        cleaned_response = cleaned_response.replace("```json", "").replace("```", "").strip()

        parsed_data = json.loads(cleaned_response)

        with open("parsed_resume.json", "w", encoding="utf-8") as f:
            json.dump(parsed_data, f, indent=2)

        return parsed_data

    except json.JSONDecodeError as e:
        print("❌ JSON parsing error:", e)
        print("Raw response from LLM:\n", response_text)
        return {}
    except Exception as e:
        print("❌ Error calling Gemini API:", e)
        return {}

# ==============================================================================
# 3. DOCX RESUME GENERATION MODULE (UPDATED)
# ==============================================================================
# This module creates a formatted DOCX file from the parsed and enhanced data.



def _apply_theme_body_font(run, size_pt=None, bold=False, underline=False, color=None):
    """Apply Calibri (Body) theme font safely to a run."""
    r = run._element

    # Ensure rPr exists
    if r.rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    else:
        rPr = r.rPr

    # Ensure rFonts exists
    if rPr.rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    else:
        rFonts = rPr.rFonts

    # Set the theme font
    rFonts.set(qn('w:asciiTheme'), 'minorHAnsi')  # Calibri (Body)
    rFonts.set(qn('w:hAnsiTheme'), 'minorHAnsi')

    # Font size
    if size_pt:
        run.font.size = Pt(size_pt)

    # Bold, underline, color
    run.bold = bold
    run.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)



def _add_section_heading(doc, text):
    """Add a section heading with correct style."""
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(text)
    _apply_theme_body_font(run, size_pt=14, bold=True, underline=True)

def _add_logo_top_right(doc, logo_path):
    """Place logo at top right of page 1 only."""
    if os.path.exists(logo_path):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(5.5)
        table.columns[1].width = Inches(1.5)
        cell_left, cell_right = table.rows[0].cells
        p = cell_right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(1.5))
        # Remove borders
        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(docx.oxml.parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'))

def _build_skill_matrix(doc, parsed_data):
    """Build skill matrix table."""
    skills_matrix = parsed_data.get("skills_matrix", [])
    if not skills_matrix and parsed_data.get("skills"):
        # Auto-generate from skills list
        skills_matrix = []
        for skill in parsed_data["skills"][:10]:
            skills_matrix.append({
                "skills": skill,
                "years_experience": "",
                "last_used": "",
                "proficiency": ""
            })

    if not skills_matrix:
        return

    # Limit to 10 rows
    skills_matrix = skills_matrix[:10]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row with purple background and white text
    hdr_cells = table.rows[0].cells
    headers = ["Area", "Years", "Latest 3 Clients Used", "Level"]
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        _apply_theme_body_font(run, size_pt=12, bold=True)
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="5A2A82"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

    # Data rows
    for skill in skills_matrix:
        row_cells = table.add_row().cells
        row_cells[0].text = str(skill.get("skills", ""))
        row_cells[1].text = str(skill.get("years_experience", ""))
        row_cells[2].text = str(skill.get("last_used", ""))
        row_cells[3].text = str(skill.get("proficiency", ""))

def create_formatted_docx(parsed_data, output_path, logo_path="logo.png"):
    """Create DOCX resume in company format."""
    doc = docx.Document()

    # Add logo top-right
    _add_logo_top_right(doc, logo_path)

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_text = "www.quadranttechnologies.com                             5020, 148th Avenue NE, Suite-250, Redmond, WA-98052"
    footer_run = footer_para.add_run(footer_text)
    _apply_theme_body_font(footer_run, size_pt=10)


    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Get name and format with initial capitals
    raw_name = parsed_data.get("name", "Name Not Found").strip()
    formatted_name = string.capwords(raw_name)  # Handles multi-word capitalization

    name_run = name_para.add_run(formatted_name)
    _apply_theme_body_font(name_run, size_pt=14, bold=True)

    

    # Contact Info
    contact_details = [
        parsed_data.get("phone", ""),
        parsed_data.get("email", ""),
        parsed_data.get("linkedin", "")
    ]
    for detail in filter(None, contact_details):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(detail)
        _apply_theme_body_font(run, size_pt=12)

    # Summary
    if parsed_data.get("summary"):
        _add_section_heading(doc, "Candidate Strengths")
        para = doc.add_paragraph(parsed_data["summary"])
        _apply_theme_body_font(para.runs[0], size_pt=12)

    # Skill Matrix
    _add_section_heading(doc, "Skill Matrix")
    _build_skill_matrix(doc, parsed_data)

    # Certifications
    if parsed_data.get("certifications"):
        _add_section_heading(doc, "Certifications")
        for cert in parsed_data["certifications"]:
            para = doc.add_paragraph(f"{cert.get('name', '')} - {cert.get('issuer', '')}", style='List Bullet')
            _apply_theme_body_font(para.runs[0], size_pt=12)

    # Education
    if parsed_data.get("education"):
        _add_section_heading(doc, "Education")
        for edu in parsed_data["education"]:
            edu_text = f"{edu.get('degree', '')} in {edu.get('major', '')} - {edu.get('university', '')}"
            para = doc.add_paragraph(edu_text, style='List Bullet')
            _apply_theme_body_font(para.runs[0], size_pt=12)

    # Experience
    if parsed_data.get("experience"):
        _add_section_heading(doc, "Professional Experience")
        for exp in parsed_data["experience"]:
            para = doc.add_paragraph(f"{exp.get('job_title', '')} - {exp.get('company', '')} ({exp.get('start_date', '')} - {exp.get('end_date', '')})")
            _apply_theme_body_font(para.runs[0], size_pt=12, bold=True)
            for ach in exp.get("Description", []):
                bullet = doc.add_paragraph(ach, style='List Bullet')
                _apply_theme_body_font(bullet.runs[0], size_pt=12)

    # Projects
    if parsed_data.get("projects"):
        _add_section_heading(doc, "Projects")
        for proj in parsed_data["projects"]:
            para = doc.add_paragraph(f"{proj.get('project_name', '')} - {proj.get('client', '')} ({proj.get('date_range', '')})")
            _apply_theme_body_font(para.runs[0], size_pt=12, bold=True)
            for c in proj.get("content", []):
                bullet = doc.add_paragraph(c, style='List Bullet')
                _apply_theme_body_font(bullet.runs[0], size_pt=12)
            if proj.get("technologies"):
                doc.add_paragraph(f"Technologies: {proj['technologies']}")
            if proj.get("environment"):
                doc.add_paragraph(f"Environment: {proj['environment']}")

    # Save file
    try:
        doc.save(output_path)
        print(f"✅ Successfully generated formatted resume at: {output_path}")
    except Exception as e:
        print(f"❌ Error saving DOCX file: {e}")


# ==============================================================================
    
if __name__ == "__main__":
    # Change this path to the PDF/DOCX you want to test

    resume_path = "YOUR_RESUME_PATH_HERE.pdf"  # Change to your resume file path
    # Load your Gemini API key from environment
    GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
    if not GEMINI_API_KEY:
        print("❌ Gemini API key not found. Please set GEMINI_API_KEY in your environment or .env file.")
        exit(1)

    # Step 1: Extract text
    extracted = extract_text(resume_path)
    if extracted:
        print("\n===== RAW EXTRACTED TEXT =====\n")
        print(extracted)
        print("\n===== END OF EXTRACTED TEXT =====\n")

        # Step 2: Parse content with LLM
        parsed_data = parse_content(extracted, GEMINI_API_KEY)

        # Step 3: Print parsed JSON
        print("\n===== PARSED RESUME JSON =====\n")
        print(json.dumps(parsed_data, indent=2))
        print("\n===== END OF PARSED RESUME JSON =====\n")
        
        #step 4: Generate formatted DOCX
        #safe_name = re.sub(r'[^A-Za-z0-9]+','-' ,'_', parsed_data['name'])
        
        safe_name = re.sub(r'[^A-Za-z0-9]+', '_', string.capwords(parsed_data['name']))

        output_path = f"YOUR_FILE_PATH\\{safe_name}_resume.docx"
        create_formatted_docx(parsed_data, output_path, logo_path="logo.png")

    else:
        print("❌ No text extracted.")

        
    
